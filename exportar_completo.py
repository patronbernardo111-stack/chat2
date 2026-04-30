#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import re
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import io

# Configuracion
EXTENSIONES_CODIGO = {".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".htm", ".css", ".scss", ".sass", ".json", ".xml", ".yaml", ".yml", ".md", ".txt", ".env", ".gitignore", ".dockerignore", "Dockerfile", "docker-compose.yml", ".sql", ".sh", ".bat", ".ps1", ".java", ".c", ".cpp", ".h", ".php", ".rb", ".go", ".rs", ".swift", ".kt"}
EXTENSIONES_IMAGENES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".ico", ".svg", ".webp", ".tiff"}
EXCLUIR_CARPETAS = {".git", "__pycache__", "node_modules", "dist", "build", ".idea", ".vscode", "venv", "env", "coverage", ".pytest_cache", ".next", "out", "target", "bin", "obj"}
MAX_ARCHIVOS = 2000  # Limite por seguridad
ANCHO_IMAGEN_CM = 10  # Ancho maximo para imagenes en Word

def limpiar_texto(texto):
    """Elimina caracteres que rompen XML"""
    if not texto:
        return ""
    # Elimina NULL bytes y caracteres de control excepto tab, newline, carriage return
    texto_limpio = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', texto)
    # Elimina BOM si existe
    if texto_limpio.startswith('\ufeff'):
        texto_limpio = texto_limpio[1:]
    return texto_limpio

def es_archivo_texto(ruta):
    """Determina si un archivo es de texto (se puede leer)"""
    ext = ruta.suffix.lower()
    if ext in EXTENSIONES_CODIGO:
        return True
    # Intentar leer primeros bytes para detectar si es binario
    try:
        with open(ruta, 'rb') as f:
            chunk = f.read(1024)
        return not b'\x00' in chunk  # Si tiene NULL bytes, es binario
    except:
        return False

def leer_contenido_seguro(ruta):
    """Lee archivo de texto manejando diferentes codificaciones"""
    for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
        try:
            with open(ruta, 'r', encoding=encoding) as f:
                contenido = f.read()
            return limpiar_texto(contenido)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "[No se pudo leer el archivo: codificacion desconocida]"

def recopilar_todos_archivos(ruta_base):
    """Recopila TODOS los archivos del proyecto"""
    ruta_base = Path(ruta_base).resolve()
    archivos = []
    
    for ruta in ruta_base.rglob("*"):
        if not ruta.is_file():
            continue
        # Excluir carpetas
        if any(excluir in ruta.parts for excluir in EXCLUIR_CARPETAS):
            continue
        
        archivos.append(ruta)
        
        if len(archivos) >= MAX_ARCHIVOS:
            print(f"   Advertencia: L?mite de {MAX_ARCHIVOS} archivos alcanzado")
            break
    
    return sorted(archivos)

def agregar_archivo_texto(doc, ruta_archivo, ruta_relativa, contenido):
    """Agrega un archivo de texto al documento"""
    # Titulo
    p = doc.add_paragraph()
    run = p.add_run(f"?? {ruta_relativa}")
    run.bold = True
    run.font.size = Pt(10)
    
    # Contenido
    if contenido and contenido.strip():
        p_codigo = doc.add_paragraph()
        run_codigo = p_codigo.add_run(contenido)
        run_codigo.font.name = "Consolas"
        run_codigo.font.size = Pt(8)
    else:
        doc.add_paragraph("[Archivo vacio]")
    
    doc.add_paragraph("?" * 70)

def agregar_imagen(doc, ruta_imagen, ruta_relativa):
    """Agrega una imagen al documento"""
    try:
        from PIL import Image
        with Image.open(ruta_imagen) as img:
            # Calcular tama?o
            ancho_pulgadas = ANCHO_IMAGEN_CM / 2.54
            doc.add_paragraph().add_run(f"??? {ruta_relativa}").bold = True
            doc.add_picture(str(ruta_imagen), width=Inches(ancho_pulgadas))
            doc.add_paragraph(f"  Dimensiones originales: {img.width} x {img.height} px")
            doc.add_paragraph()
    except Exception as e:
        doc.add_paragraph(f"?? {ruta_relativa} - No se pudo incrustar la imagen: {str(e)[:100]}")
        doc.add_paragraph()

def agregar_archivo_binario(doc, ruta_archivo, ruta_relativa):
    """Agrega informacion sobre archivos binarios que no se pueden incrustar"""
    tamanio = ruta_archivo.stat().st_size
    if tamanio < 1024:
        tam_str = f"{tamanio} bytes"
    elif tamanio < 1024*1024:
        tam_str = f"{tamanio/1024:.1f} KB"
    else:
        tam_str = f"{tamanio/(1024*1024):.1f} MB"
    
    p = doc.add_paragraph()
    run = p.add_run(f"?? {ruta_relativa}")
    run.bold = True
    run.font.size = Pt(9)
    doc.add_paragraph(f"  Tipo: Archivo binario | Tama?o: {tam_str}")
    doc.add_paragraph()

def main():
    print("=" * 60)
    print("EXPORTADOR COMPLETO DE PROYECTO A WORD")
    print("CODIGO + IMAGENES + TODOS LOS ARCHIVOS")
    print("=" * 60)
    print()
    
    ruta_input = input("Ingresa la ruta completa de tu proyecto: ").strip()
    
    if not ruta_input:
        print("? No ingresaste ninguna ruta")
        return
    
    ruta_proyecto = Path(ruta_input)
    if not ruta_proyecto.exists():
        print(f"? La ruta '{ruta_input}' no existe")
        return
    
    if not ruta_proyecto.is_dir():
        print(f"? La ruta no es una carpeta")
        return
    
    print(f"\n?? Proyecto: {ruta_proyecto.name}")
    print(f"?? Ruta: {ruta_proyecto}")
    print("\n?? Recopilando TODOS los archivos...")
    
    archivos = recopilar_todos_archivos(ruta_proyecto)
    total = len(archivos)
    
    if total == 0:
        print("? No se encontraron archivos")
        return
    
    # Clasificar
    archivos_texto = []
    archivos_imagen = []
    archivos_otros = []
    
    for arch in archivos:
        ext = arch.suffix.lower()
        if ext in EXTENSIONES_IMAGENES:
            archivos_imagen.append(arch)
        elif es_archivo_texto(arch):
            archivos_texto.append(arch)
        else:
            archivos_otros.append(arch)
    
    print(f"\n? Archivos encontrados:")
    print(f"   ?? C?digo/Texto: {len(archivos_texto)}")
    print(f"   ??? Imagenes/Icons: {len(archivos_imagen)}")
    print(f"   ?? Otros archivos: {len(archivos_otros)}")
    print(f"   ?? TOTAL: {total} archivos")
    
    if total > 1000:
        print(f"\n?? ADVERTENCIA: El proyecto tiene {total} archivos")
        print("   El documento Word resultante podria ser EXTREMADAMENTE GRANDE")
        respuesta = input("   ?Deseas continuar igual? (s/n): ").strip().lower()
        if respuesta != 's':
            print("   Cancelado por seguridad")
            return
    
    print("\n?? Generando documento Word...")
    
    # Crear documento
    doc = Document()
    
    # Portada
    titulo = doc.add_heading(f"PROYECTO COMPLETO: {ruta_proyecto.name}", 0)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    doc.add_paragraph(f"Fecha de exportacion: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    doc.add_paragraph(f"Ruta original: {ruta_proyecto}")
    doc.add_paragraph()
    
    # Resumen
    doc.add_heading("Resumen del proyecto", level=1)
    doc.add_paragraph(f"? Total de archivos: {total}")
    doc.add_paragraph(f"? Archivos de codigo: {len(archivos_texto)}")
    doc.add_paragraph(f"? Imagenes e iconos: {len(archivos_imagen)}")
    doc.add_paragraph(f"? Otros archivos: {len(archivos_otros)}")
    doc.add_page_break()
    
    # Indice de archivos
    doc.add_heading("Listado completo de archivos", level=1)
    for arch in archivos[:500]:  # Limite para no saturar
        try:
            rel = arch.relative_to(ruta_proyecto)
            doc.add_paragraph(str(rel), style='List Bullet')
        except:
            pass
    if total > 500:
        doc.add_paragraph(f"... y {total - 500} archivos mas")
    doc.add_page_break()
    
    # Procesar archivos de texto/codigo
    if archivos_texto:
        doc.add_heading("?? CODIGO FUENTE Y ARCHIVOS DE TEXTO", level=1)
        for idx, arch in enumerate(archivos_texto, 1):
            try:
                rel = arch.relative_to(ruta_proyecto)
                print(f"   Procesando texto {idx}/{len(archivos_texto)}: {rel}")
                contenido = leer_contenido_seguro(arch)
                agregar_archivo_texto(doc, arch, rel, contenido)
            except Exception as e:
                print(f"   Error: {e}")
                doc.add_paragraph(f"?? Error procesando {rel}: {str(e)[:100]}")
    
    # Procesar imagenes
    if archivos_imagen:
        doc.add_page_break()
        doc.add_heading("??? IMAGENES E ICONOS", level=1)
        for idx, arch in enumerate(archivos_imagen, 1):
            try:
                rel = arch.relative_to(ruta_proyecto)
                print(f"   Procesando imagen {idx}/{len(archivos_imagen)}: {rel}")
                agregar_imagen(doc, arch, rel)
            except Exception as e:
                print(f"   Error: {e}")
                doc.add_paragraph(f"?? Error con imagen {rel}: {str(e)[:100]}")
    
    # Procesar otros archivos
    if archivos_otros:
        doc.add_page_break()
        doc.add_heading("?? OTROS ARCHIVOS", level=1)
        for idx, arch in enumerate(archivos_otros, 1):
            try:
                rel = arch.relative_to(ruta_proyecto)
                print(f"   Procesando otro {idx}/{len(archivos_otros)}: {rel}")
                agregar_archivo_binario(doc, arch, rel)
            except Exception as e:
                print(f"   Error: {e}")
    
    # Guardar
    nombre_salida = f"proyecto_completo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    ruta_salida = Path.cwd() / nombre_salida
    doc.save(ruta_salida)
    
    tamanio_mb = ruta_salida.stat().st_size / (1024 * 1024)
    
    print("\n" + "=" * 60)
    print("? EXPORTACION COMPLETADA CON EXITO!")
    print(f"?? Archivo guardado: {ruta_salida}")
    print(f"?? Tama?o del archivo: {tamanio_mb:.2f} MB")
    print(f"?? Total archivos incluidos: {total}")
    print("=" * 60)
    
    if tamanio_mb > 100:
        print("\n?? El archivo es muy grande (>100 MB). Puede tardar en abrirse.")
    else:
        print("\n? Puedes abrir el archivo con Microsoft Word o LibreOffice")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n? Exportacion cancelada por el usuario")
    except Exception as e:
        print(f"\n? Error inesperado: {e}")
        print("\nSi el error persiste, es posible que tu proyecto sea demasiado grande.")
